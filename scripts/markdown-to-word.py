#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script chuyển đổi Markdown sang Word (.docx)
"""

import re
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Đang cài đặt python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

def parse_markdown(markdown_content):
    """Parse markdown content thành các elements"""
    lines = markdown_content.split('\n')
    elements = []
    in_code_block = False
    code_block_content = []
    code_language = ''
    
    for line in lines:
        # Xử lý code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_language = line.strip()[3:].strip()
                code_block_content = []
            else:
                in_code_block = False
                if code_block_content:
                    elements.append({
                        'type': 'code',
                        'language': code_language,
                        'content': '\n'.join(code_block_content)
                    })
                code_block_content = []
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue
        
        # Xử lý headings
        if line.startswith('# '):
            elements.append({'type': 'h1', 'text': line[2:].strip()})
        elif line.startswith('## '):
            elements.append({'type': 'h2', 'text': line[3:].strip()})
        elif line.startswith('### '):
            elements.append({'type': 'h3', 'text': line[4:].strip()})
        elif line.startswith('#### '):
            elements.append({'type': 'h4', 'text': line[5:].strip()})
        elif line.startswith('##### '):
            elements.append({'type': 'h5', 'text': line[6:].strip()})
        elif line.startswith('###### '):
            elements.append({'type': 'h6', 'text': line[7:].strip()})
        # Xử lý horizontal rule
        elif line.strip() == '---' or line.strip() == '***':
            elements.append({'type': 'hr'})
        # Xử lý list items
        elif line.strip().startswith('- '):
            text = line.strip()[2:].strip()
            elements.append({'type': 'bullet', 'text': text})
        # Xử lý empty lines
        elif line.strip() == '':
            elements.append({'type': 'empty'})
        # Xử lý text thường
        elif line.strip():
            elements.append({'type': 'text', 'text': line.strip()})
    
    return elements

def parse_text_runs(text):
    """Parse text để xử lý bold và italic"""
    runs = []
    current_pos = 0
    
    # Tìm tất cả các markdown formatting
    patterns = [
        (r'\*\*(.+?)\*\*', 'bold'),  # Bold
        (r'\*(.+?)\*', 'italic'),      # Italic
        (r'`(.+?)`', 'code'),         # Inline code
    ]
    
    matches = []
    for pattern, style in patterns:
        for match in re.finditer(pattern, text):
            matches.append({
                'start': match.start(),
                'end': match.end(),
                'text': match.group(1),
                'style': style
            })
    
    # Sắp xếp theo vị trí
    matches.sort(key=lambda x: x['start'])
    
    # Loại bỏ overlaps
    filtered_matches = []
    for match in matches:
        overlap = False
        for existing in filtered_matches:
            if (match['start'] < existing['end'] and match['end'] > existing['start']):
                overlap = True
                break
        if not overlap:
            filtered_matches.append(match)
    
    # Tạo runs
    last_pos = 0
    for match in filtered_matches:
        if match['start'] > last_pos:
            runs.append({'text': text[last_pos:match['start']], 'style': None})
        
        runs.append({
            'text': match['text'],
            'style': match['style']
        })
        
        last_pos = match['end']
    
    if last_pos < len(text):
        runs.append({'text': text[last_pos:], 'style': None})
    
    return runs if runs else [{'text': text, 'style': None}]

def create_word_document(elements, output_path):
    """Tạo Word document từ các elements"""
    doc = Document()
    
    # Thiết lập font mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Tạo style cho code
    try:
        code_style = doc.styles.add_style('Code', 1)  # 1 = paragraph style
    except:
        code_style = doc.styles['Code']
    
    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = Pt(10)
    
    for element in elements:
        if element['type'] == 'h1':
            p = doc.add_heading(element['text'], level=1)
            p.paragraph_format.space_after = Pt(12)
        elif element['type'] == 'h2':
            p = doc.add_heading(element['text'], level=2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(10)
        elif element['type'] == 'h3':
            p = doc.add_heading(element['text'], level=3)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(8)
        elif element['type'] == 'h4':
            p = doc.add_heading(element['text'], level=4)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(6)
        elif element['type'] == 'h5':
            p = doc.add_heading(element['text'], level=5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
        elif element['type'] == 'h6':
            p = doc.add_heading(element['text'], level=6)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif element['type'] == 'bullet':
            p = doc.add_paragraph(element['text'], style='List Bullet')
            p.paragraph_format.space_after = Pt(6)
        elif element['type'] == 'code':
            p = doc.add_paragraph(element['content'], style='Code')
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.25)
        elif element['type'] == 'hr':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p_format = p.paragraph_format
            p_format.border_bottom.color.rgb = RGBColor(0, 0, 0)
            p_format.border_bottom.width = Pt(0.5)
        elif element['type'] == 'text':
            runs = parse_text_runs(element['text'])
            p = doc.add_paragraph()
            for run_data in runs:
                run = p.add_run(run_data['text'])
                if run_data['style'] == 'bold':
                    run.bold = True
                elif run_data['style'] == 'italic':
                    run.italic = True
                elif run_data['style'] == 'code':
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(6)
        elif element['type'] == 'empty':
            doc.add_paragraph()
    
    # Lưu file
    doc.save(output_path)
    print(f"✅ Đã xuất thành công file Word tại: {output_path}")

def main():
    # Đường dẫn file
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    markdown_path = project_root / 'KIEN-TRUC-HE-THONG.md'
    output_path = project_root / 'KIEN-TRUC-HE-THONG.docx'
    
    if not markdown_path.exists():
        print(f"❌ Không tìm thấy file: {markdown_path}")
        sys.exit(1)
    
    print(f"Đang đọc file markdown: {markdown_path}")
    with open(markdown_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    print("Đang parse markdown...")
    elements = parse_markdown(markdown_content)
    
    print("Đang tạo Word document...")
    create_word_document(elements, output_path)
    
    print("Hoàn thành!")

if __name__ == '__main__':
    main()
